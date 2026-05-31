import os
import shutil
import docx.shared
from docx import Document
from datetime import datetime

class ExportManager:
    """Maneja la exportación de datos del proyecto (Diario, Excel, etc) puro, sin dependencias GUI."""

    @staticmethod
    def export_diary(entries, project_name, export_path):
        """Exporta el diario estructurado (lista de diccionarios) a un documento Word."""
        doc = Document()
        doc.add_heading(f"Diario de codificación - {project_name}", level=1)

        if not entries:
            doc.add_paragraph("(Diario vacío)")
        else:
            for entry in entries:
                # 1. Parsear y formatear la fecha
                date_str = entry.get("date", "")
                try:
                    dt = datetime.fromisoformat(date_str)
                    date_formatted = dt.strftime("%d/%m/%Y a las %H:%M")
                except ValueError:
                    date_formatted = date_str  # Fallback por si la fecha está corrupta o en otro formato

                author = entry.get("author", "Desconocido")
                message = entry.get("message", "")

                # 2. Agregar encabezado de la entrada (Autor y Fecha)
                p_header = doc.add_paragraph()
                run = p_header.add_run(f"👤 {author} - {date_formatted}")
                run.bold = True

                # 3. Agregar el cuerpo del mensaje
                doc.add_paragraph(message)
                
                # 4. Agregar línea separadora sutil
                p_divider = doc.add_paragraph()
                run_divider = p_divider.add_run("_" * 50)
                run_divider.font.color.rgb = docx.shared.RGBColor(180, 180, 180) # Gris claro

        doc.save(export_path)

    @staticmethod
    def export_code_tree(rows, export_path):
        """Exporta el libro de códigos (jerarquía) a Excel."""
        import openpyxl
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Libro de códigos"
        
        ExportManager._write_codebook_sheet(ws, rows, title_text="Libro de códigos", frequency_label="Frecuencia")
        wb.save(export_path)

    @staticmethod
    def export_code_fragments(selected_rows, fragment_rows, export_path):
        """Exporta el libro de códigos y sus fragmentos asociados a Excel."""
        import openpyxl

        wb = openpyxl.Workbook()
        
        # Hoja 1: Libro
        ws_codes = wb.active
        ws_codes.title = "Libro de códigos"
        ExportManager._write_codebook_sheet(ws_codes, selected_rows, title_text="Libro de códigos", frequency_label="Fragmentos")
        
        # Hoja 2: Fragmentos
        ws_fragments = wb.create_sheet("Fragmentos")
        ExportManager._write_fragments_sheet(ws_fragments, fragment_rows)
        
        wb.save(export_path)

    @staticmethod
    def _write_codebook_sheet(worksheet, rows, title_text="Libro de códigos", frequency_label="Frecuencia"):
        from openpyxl.styles import Font, PatternFill, Alignment

        header = [title_text, "", "", "", "Memo", frequency_label]
        worksheet.append(header)

        header_fill = PatternFill(start_color="5d9bd3", end_color="5d9bd3", fill_type="solid")
        data_fill = PatternFill(start_color="f6f8fb", end_color="f6f8fb", fill_type="solid")
        bold = Font(bold=True)
        memo_col = 5
        freq_col = 6

        for col_idx in range(1, len(header) + 1):
            cell = worksheet.cell(row=1, column=col_idx)
            cell.font = bold
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="center")

        for row_data in rows:
            level = row_data.get("level", 0)
            total_cols = max(freq_col, 1 + level)
            row = ["" for _ in range(total_cols)]
            name_col = 1 + level
            row[name_col - 1] = row_data.get("name", "")
            
            if memo_col - 1 >= len(row):
                row.extend([""] * (memo_col - len(row)))
            row[memo_col - 1] = row_data.get("memo", "")
            
            if freq_col - 1 >= len(row):
                row.extend([""] * (freq_col - len(row)))
            row[freq_col - 1] = row_data.get("freq", "")
            
            worksheet.append(row)
            current_row = worksheet.max_row
            for col_idx in range(1, freq_col + 1):
                cell = worksheet.cell(row=current_row, column=col_idx)
                cell.fill = data_fill
                cell.alignment = Alignment(vertical="top", wrap_text=True)

        for col_name in ("A", "B", "C", "D"):
            worksheet.column_dimensions[col_name].width = 36
        worksheet.column_dimensions["E"].width = 48
        worksheet.column_dimensions["F"].width = 14

        for row in worksheet.iter_rows(min_row=1, max_row=worksheet.max_row, min_col=1, max_col=freq_col):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    @staticmethod
    def _write_fragments_sheet(worksheet, fragment_rows):
        from openpyxl.styles import Font, PatternFill, Alignment

        header = ["Código", "Documento", "Fragmento", "Memo"]
        worksheet.append(header)

        header_fill = PatternFill(start_color="5d9bd3", end_color="5d9bd3", fill_type="solid")
        data_fill = PatternFill(start_color="f6f8fb", end_color="f6f8fb", fill_type="solid")
        bold = Font(bold=True)

        for col_idx in range(1, len(header) + 1):
            cell = worksheet.cell(row=1, column=col_idx)
            cell.font = bold
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="center")

        for fragment in fragment_rows:
            worksheet.append([
                fragment.get("code_name", ""),
                fragment.get("document", ""),
                fragment.get("text", ""),
                fragment.get("memo", ""),
            ])
            current_row = worksheet.max_row
            for col_idx in range(1, len(header) + 1):
                cell = worksheet.cell(row=current_row, column=col_idx)
                cell.fill = data_fill
                cell.alignment = Alignment(vertical="top", wrap_text=True)

        worksheet.column_dimensions["A"].width = 28
        worksheet.column_dimensions["B"].width = 28
        worksheet.column_dimensions["C"].width = 100
        worksheet.column_dimensions["D"].width = 48

    @staticmethod
    def export_project_to_rqa(project_path: str, export_path: str) -> str:
        """
        Comprime la carpeta entera del proyecto en un archivo .rqa.
        Si la export_path no termina en .rqa, se le añade.
        """
        if not export_path.lower().endswith(".rqa"):
            export_path += ".rqa"
            
        base_name = export_path[:-4] if export_path.lower().endswith(".rqa") else export_path
        
        zip_path = shutil.make_archive(base_name, 'zip', project_path)
        
        if os.path.exists(export_path):
            os.remove(export_path) 
            
        shutil.move(zip_path, export_path)
        
        return export_path

    @staticmethod
    def export_exchange_to_rex(project, selected_docs, selected_codes, options, export_path):
        """
        Exporta los documentos, códigos, temas y memos seleccionados a un archivo
        de intercambio con formato .rex.

        Args:
            project (Project): El proyecto actual de RaizQA del cual extraer los datos.
            selected_docs (list): Lista de nombres de documentos a exportar.
            selected_codes (list): Lista de nombres de códigos a exportar.
            options (dict): Diccionario con opciones extra (como include_memos y code_themes).
            export_path (str): Ruta donde se guardará el archivo .rex generado.
        """
        import json
        import zipfile
        import tempfile
        
        include_memos = options.get("include_memos", False)
        
        if not export_path.lower().endswith(".rex"):
            export_path += ".rex"
            
        with tempfile.TemporaryDirectory() as temp_dir:
            docs_dir = os.path.join(temp_dir, "documentos")
            os.makedirs(docs_dir)
            
            # Copy documents
            for doc in selected_docs:
                src_path = os.path.join(project.documents_path, doc)
                if os.path.exists(src_path):
                    shutil.copy2(src_path, os.path.join(docs_dir, doc))
                    
            # Filter codes_dict
            filtered_codes = {}
            for code in selected_codes:
                if code in project.codes_dict:
                    code_data = project.codes_dict[code].copy()
                    code_data["fragments"] = {}
                    for doc, frags in project.codes_dict[code].get("fragments", {}).items():
                        if doc in selected_docs:
                            code_data["fragments"][doc] = frags
                    filtered_codes[code] = code_data
                    
            # Filter themes using the ones passed from UI if available
            code_themes = options.get("code_themes", [])
            filtered_themes = {}
            if code_themes:
                for theme in code_themes:
                    theme_name = theme.get("name", "Tema sin nombre")
                    theme_codes = [c for c in theme.get("codes", []) if c in selected_codes]
                    if theme_codes:
                        filtered_themes[theme_name] = {
                            "memo": "",  # Or you could fetch the memo from project.themes_dict if it exists
                            "codes": theme_codes
                        }
            else:
                for theme_name, theme_data in project.themes_dict.items():
                    theme_codes = [c for c in theme_data.get("codes", []) if c in selected_codes]
                    if theme_codes:
                        filtered_themes[theme_name] = {
                            "memo": theme_data.get("memo", ""),
                            "codes": theme_codes
                        }
                    
            # Filter memos_dict
            filtered_memos = {}
            if include_memos:
                for memo_id, memo_text in project.memos_dict.items():
                    target = memo_id
                    if target in selected_docs or target in selected_codes:
                        filtered_memos[memo_id] = memo_text
                        
            project_data = {
                "codes_dict": filtered_codes,
                "themes_dict": filtered_themes,
                "memos_dict": filtered_memos,
            }
            
            with open(os.path.join(temp_dir, "project_data.json"), "w", encoding="utf-8") as f:
                json.dump(project_data, f, indent=4, ensure_ascii=False)
                
            metadata = {
                "name": project.name,
                "documents": selected_docs,
                "is_exchange": True
            }
            with open(os.path.join(temp_dir, "metadata.json"), "w", encoding="utf-8") as f:
                json.dump(metadata, f, indent=4, ensure_ascii=False)
                
            base_name = export_path[:-4]
            zip_path = shutil.make_archive(base_name, 'zip', temp_dir)
            
            if os.path.exists(export_path):
                os.remove(export_path)
            shutil.move(zip_path, export_path)
            
        return export_path
